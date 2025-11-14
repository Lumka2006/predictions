import os, pandas as pd, numpy as np, matplotlib.pyplot as plt, seaborn as sns, joblib, warnings, statsmodels.api as sm
from sklearn.preprocessing import LabelEncoder
from sklearn.feature_selection import SelectKBest, f_classif
from sklearn.ensemble import RandomForestClassifier
from sklearn.impute import SimpleImputer
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.metrics import accuracy_score, confusion_matrix, classification_report, multilabel_confusion_matrix
from scipy import stats
from statsmodels.formula.api import ols
from docx import Document
warnings.filterwarnings("ignore")

# Load Data
data_file_path = "data.xlsx"; target_column = "Target"
if not os.path.exists(data_file_path): raise FileNotFoundError(f"{data_file_path} not found!")
dataframe = pd.read_excel(data_file_path)
if target_column not in dataframe: raise ValueError(f"Missing target column '{target_column}'!")
print(f"✅ Data loaded: {dataframe.shape}")

# Encode Categorical Variables
label_encoders = {}
for column in dataframe.select_dtypes('object'):
     label_encoders[column]=LabelEncoder(); dataframe[column]=label_encoders[column].fit_transform(dataframe[column].astype(str)); 
if column==target_column: print("Target encoding mapping:", dict(zip(label_encoders[column].classes_, label_encoders[column].transform(label_encoders[column].classes_))))

# Split Features and Target
features, target = dataframe.drop(target_column, axis=1), dataframe[target_column]

# Feature Importance
anova_selector = SelectKBest(f_classif, k='all').fit(features, target); 
random_forest_model = RandomForestClassifier(random_state=42).fit(features, target)
feature_importance_dataframe = pd.DataFrame({"ANOVA_Score": anova_selector.scores_, "RandomForest_Importance": random_forest_model.feature_importances_}, index=features.columns)
feature_importance_dataframe = feature_importance_dataframe.div(feature_importance_dataframe.max()); feature_importance_dataframe["Overall_Importance"] = feature_importance_dataframe.mean(axis=1); feature_importance_dataframe = feature_importance_dataframe.sort_values("Overall_Importance", ascending=False)

# Save feature importance to Excel
with pd.ExcelWriter("FeatureImportance.xlsx") as excel_writer: 
    pd.Series(anova_selector.scores_, index=features.columns).sort_values(ascending=False).to_excel(excel_writer, "ANOVA_Scores"); 
    pd.Series(random_forest_model.feature_importances_, index=features.columns).sort_values(ascending=False).to_excel(excel_writer, "RF_Importance"); 
    feature_importance_dataframe.to_excel(excel_writer, "Combined_Importance")

# Select Top 15 Features
top_15_features = feature_importance_dataframe.head(15).index.tolist(); 
dataframe[top_15_features + [target_column]].to_excel("SelectedFeatures.xlsx", index=False)

# Plot Overall Importance
feature_importance_dataframe["Overall_Importance"].plot.barh(color="skyblue"); plt.gca().invert_yaxis(); plt.show()

# Data Cleaning
dataframe = pd.read_excel("SelectedFeatures.xlsx"); 
dataframe.columns = dataframe.columns.str.strip().str.replace(r"[() ]", "_", regex=True)
numerical_columns, categorical_columns = dataframe.select_dtypes(np.number).columns, dataframe.select_dtypes(exclude=np.number).columns
dataframe[numerical_columns] = SimpleImputer(strategy="median").fit_transform(dataframe[numerical_columns]); 
dataframe[categorical_columns] = SimpleImputer(strategy="most_frequent").fit_transform(dataframe[categorical_columns]) if len(categorical_columns) else dataframe[categorical_columns]
print("✅ Data cleaned")

# Detect Outliers
outlier_rows = dataframe[(np.abs(stats.zscore(dataframe[numerical_columns])) > 3).any(axis=1)]; 
outlier_rows.to_excel("outliers.xlsx", index=False); print(f"📊 {len(outlier_rows)} outliers saved")

# Visualizations
sns.set_style("whitegrid")
for column in dataframe.columns:
    if np.issubdtype(dataframe[column].dtype, np.number): 
        plt.hist(dataframe[column], bins=15, color='aqua', edgecolor='black', alpha=0.7); plt.gcf().text(0.75,0.7,f'Mean:{dataframe[column].mean():.2f}\nMedian:{dataframe[column].median():.2f}\nMode:{dataframe[column].mode()[0]:.2f}', fontsize=9,bbox=dict(facecolor='white', alpha=0.6))
    else: 
        sns.countplot(x=column, data=dataframe, palette='Set3', order=dataframe[column].value_counts().index); 
        plt.xticks(rotation=45)
    plt.title(column); plt.tight_layout(); plt.show()

# ANOVA Test
with pd.ExcelWriter("anova_results.xlsx") as excel_writer: [sm.stats.anova_lm(ols(f'Q("{col}") ~ C({target_column})', data=dataframe).fit(), typ=2).to_excel(excel_writer, sheet_name=col[:31]) for col in numerical_columns[:15]]
print("✅ ANOVA results saved")

# Train/Test Split
features, target = dataframe.drop(target_column, axis=1), dataframe[target_column]
features_train, features_test, target_train, target_test = train_test_split(features, target, test_size=0.1, random_state=42)
features_train.assign(Target=target_train).to_excel("train_data.xlsx", index=False); features_test.assign(Target=target_test).to_excel("test_data.xlsx", index=False)

# Random Forest with GridSearchCV
random_forest_param_grid = {
    'n_estimators':[300,400,500],
    'max_depth':[None,10,20],
    'min_samples_split':[2,5],
    'min_samples_leaf':[1,2],
    'max_features':['sqrt','log2']}
grid_search = GridSearchCV(estimator=RandomForestClassifier(random_state=42), param_grid=random_forest_param_grid, cv=3, n_jobs=-1, scoring='accuracy', verbose=1); 
grid_search.fit(features_train, target_train)
print("✅ Best parameters found:", grid_search.best_params_)

optimized_random_forest_model = grid_search.best_estimator_; 
predicted_target_test = optimized_random_forest_model.predict(features_test)
pd.concat([features_test, pd.Series(predicted_target_test, name=target_column)], axis=1).to_excel("predicted_results.xlsx", index=False)

# Evaluation Report
target_labels = optimized_random_forest_model.classes_; 
multi_label_conf_matrix = multilabel_confusion_matrix(target_test, predicted_target_test, labels=target_labels); 
overall_conf_matrix = confusion_matrix(target_test, predicted_target_test, labels=target_labels)

document = Document(); document.add_heading("🎓 Model Evaluation Report", 1); 
document.add_heading("Overall Performance", 2)
classification_report_dict = classification_report(target_test, predicted_target_test, output_dict=True); 
overall_accuracy = accuracy_score(target_test, predicted_target_test)

# table (3 rows and 5 columns)
table = document.add_table(rows=3, cols=5)
table.style = "Light Grid Accent 1"

# Set header row
headers = ["Metric", "Accuracy", "Precision", "Recall", "F1-Score"]
for col_idx, header in enumerate(headers):
    table.cell(0, col_idx).text = header

# Fill Weighted Average row
table.cell(1, 0).text = "Weighted Average"
table.cell(1, 1).text = f"{overall_accuracy:.4f}"
table.cell(1, 2).text = f"{classification_report_dict['weighted avg']['precision']:.4f}"
table.cell(1, 3).text = f"{classification_report_dict['weighted avg']['recall']:.4f}"
table.cell(1, 4).text = f"{classification_report_dict['weighted avg']['f1-score']:.4f}"

# Fill Macro Average row
table.cell(2, 0).text = "Macro Average"
table.cell(2, 1).text = "-"
table.cell(2, 2).text = f"{classification_report_dict['macro avg']['precision']:.4f}"
table.cell(2, 3).text = f"{classification_report_dict['macro avg']['recall']:.4f}"
table.cell(2, 4).text = f"{classification_report_dict['macro avg']['f1-score']:.4f}"

# Overall Confusion Matrix
document.add_heading("Overall Confusion Matrix", level=2)

# Create table with extra row/column for labels
rows, cols = overall_conf_matrix.shape
cm_table = document.add_table(rows=rows + 1, cols=cols + 1)
cm_table.style = "Light Grid Accent 1"

# Set top-left header
cm_table.cell(0, 0).text = "Actual / Predicted"

# Set column headers (Predicted labels)
for j, label in enumerate(target_labels):
    cm_table.cell(0, j + 1).text = str(label)

# Fill table rows (Actual labels and counts)
for i, label in enumerate(target_labels):
    cm_table.cell(i + 1, 0).text = str(label)  # Actual label
    for j in range(cols):
        cm_table.cell(i + 1, j + 1).text = str(overall_conf_matrix[i, j])

# Per-Class Metrics
document.add_heading("Per-Class Details", level=2)

for idx, label in enumerate(target_labels):
    # Compute confusion matrix values
    tn, fp, fn, tp = multi_label_conf_matrix[idx].ravel()
    total = tp + tn + fp + fn
    acc = (tp + tn) / total
    class_err = 1 - acc
    prec = tp / (tp + fp) if tp + fp else 0
    rec = tp / (tp + fn) if tp + fn else 0
    f1 = 2 * prec * rec / (prec + rec) if prec + rec else 0

    # Class heading
    document.add_heading(f"🔹 Class: {label}", level=3)

    # Formulas list
    formulas = [
        "Accuracy = (TP+TN)/(TP+TN+FP+FN)",
        "Error Rate = 1-Accuracy",
        "Precision = TP/(TP+FP)",
        "Recall = TP/(TP+FN)",
        "Specificity = TN/(TN+FP)",
        "F1-Score = 2*(Precision*Recall)/(Precision+Recall)"
    ]
    for formula in formulas:
        document.add_paragraph(f"• {formula}", style="List Bullet")

    # Metric values summary
    document.add_paragraph(
        f"Acc={acc:.4f}, Err={class_err:.4f}, Prec={prec:.4f}, Rec={rec:.4f}, F1={f1:.4f}, TP={tp}, TN={tn}, FP={fp}, FN={fn}"
    )
    # Class confusion matrix
    class_conf_table = document.add_table(rows=3, cols=3)
    class_conf_table.style = "Light Grid Accent 1"
    class_conf_table.cell(0, 1).text = "Predicted: Positive"
    class_conf_table.cell(0, 2).text = "Predicted: Negative"
    class_conf_table.cell(1, 0).text = "Actual: Positive"
    class_conf_table.cell(2, 0).text = "Actual: Negative"
    class_conf_table.cell(1, 1).text = str(tp)
    class_conf_table.cell(1, 2).text = str(fn)
    class_conf_table.cell(2, 1).text = str(fp)
    class_conf_table.cell(2, 2).text = str(tn)

# Save Report & Model
document.save("model_evaluation.docx"); joblib.dump(optimized_random_forest_model, "optimized_rf_model.pkl"); print("✅ Report & Model saved")
